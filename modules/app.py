# app.py
# Aplicación Web Streamlit - Forecast 5+7 No Lineal Dinámico
# Proyecto de gastos mineros: carga cualquier Excel con estructura de gastos/presupuesto similar.

from __future__ import annotations

import datetime
import io
import re
import unicodedata
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st

# ReportLab para PDF
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, KeepTogether
    )
    from reportlab.platypus.flowables import HRFlowable
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# python-docx para Word
try:
    import docx as _docx_module
    PYTHON_DOCX_OK = True
except ImportError:
    PYTHON_DOCX_OK = False

st.set_page_config(
    page_title="Forecast 5+7 No Lineal Dinámico",
    page_icon="⛏️",
    layout="wide",
    initial_sidebar_state="expanded",
)

MONTH_ORDER = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
MONTH_ALIASES = {
    "jan": "Jan", "january": "Jan", "ene": "Jan", "enero": "Jan", "01": "Jan", "1": "Jan",
    "feb": "Feb", "february": "Feb", "febrero": "Feb", "02": "Feb", "2": "Feb",
    "mar": "Mar", "march": "Mar", "marzo": "Mar", "03": "Mar", "3": "Mar",
    "apr": "Apr", "april": "Apr", "abr": "Apr", "abril": "Apr", "04": "Apr", "4": "Apr",
    "may": "May", "mayo": "May", "05": "May", "5": "May",
    "jun": "Jun", "june": "Jun", "junio": "Jun", "06": "Jun", "6": "Jun",
    "jul": "Jul", "july": "Jul", "julio": "Jul", "07": "Jul", "7": "Jul",
    "aug": "Aug", "august": "Aug", "ago": "Aug", "agosto": "Aug", "08": "Aug", "8": "Aug",
    "sep": "Sep", "sept": "Sep", "september": "Sep", "septiembre": "Sep", "09": "Sep", "9": "Sep",
    "oct": "Oct", "october": "Oct", "octubre": "Oct", "10": "Oct",
    "nov": "Nov", "november": "Nov", "noviembre": "Nov", "11": "Nov",
    "dec": "Dec", "december": "Dec", "dic": "Dec", "diciembre": "Dec", "12": "Dec",
}

DEFAULT_CONTEXT = {
    "Labor": {"sens": 0.35, "mom": 0.10, "curve": 0.04, "min": 0.88, "max": 1.18, "why": "gasto estructural y relativamente estable"},
    "Fuel": {"sens": 0.75, "mom": 0.35, "curve": 0.18, "min": 0.70, "max": 1.65, "why": "volátil por precio, producción y flota"},
    "Contractors": {"sens": 0.65, "mom": 0.28, "curve": 0.16, "min": 0.72, "max": 1.55, "why": "depende del avance operacional y contratos"},
    "Maintenance": {"sens": 0.55, "mom": 0.30, "curve": 0.22, "min": 0.65, "max": 1.75, "why": "asociado a mantenciones y detenciones"},
    "Power": {"sens": 0.45, "mom": 0.20, "curve": 0.08, "min": 0.80, "max": 1.35, "why": "consumo energético operacional"},
    "Spare Parts": {"sens": 0.60, "mom": 0.25, "curve": 0.16, "min": 0.70, "max": 1.60, "why": "repuestos asociados a continuidad operacional"},
    "S&C": {"sens": 0.50, "mom": 0.20, "curve": 0.10, "min": 0.78, "max": 1.40, "why": "servicios y consumibles"},
    "Rehandling": {"sens": 0.60, "mom": 0.25, "curve": 0.12, "min": 0.72, "max": 1.50, "why": "movimiento adicional de material"},
    "Other": {"sens": 0.50, "mom": 0.20, "curve": 0.10, "min": 0.75, "max": 1.45, "why": "categoría general"},
}

POSSIBLE_CLASSIF = ["classif", "clasificacion", "clasificación", "naturaleza", "nature", "categoria", "categoría", "tipo gasto", "expense type"]
POSSIBLE_CC = ["cc", "centro costo", "centro de costo", "cost center", "resp", "responsable", "ceco"]
POSSIBLE_BUDGET_FY = ["budget fy", "budget_fy", "fy25", "fy26", "fy27", "fy28", "fy29", "fy30", "budget total", "presupuesto fy", "presupuesto anual"]
POSSIBLE_FORECAST_FY = ["forecast fy", "forecast_fy", "forecast actual", "proyeccion fy", "proyección fy"]


# ─────────────────────────────────────────────
# FUNCIONES UTILITARIAS
# ─────────────────────────────────────────────

def norm_txt(x: object) -> str:
    s = str(x).strip().lower()
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"\s+", " ", s)
    return s


def find_month_in_col(col: object) -> Optional[str]:
    raw = str(col).strip()
    n = norm_txt(raw)
    if n.startswith("budget_"):
        n = n.replace("budget_", "")
    tokens = re.split(r"[^a-zA-Z0-9]+", n)
    tokens = [t for t in tokens if t]
    for t in tokens:
        if t in MONTH_ALIASES:
            return MONTH_ALIASES[t]
    for alias, std in MONTH_ALIASES.items():
        if re.search(rf"(^|[^a-z0-9]){re.escape(alias)}([^a-z0-9]|$)", n):
            return std
        if re.match(rf"^{re.escape(alias)}\d{{2,4}}$", n):
            return std
    m = re.search(r"20\d{2}[-_/ ](0?[1-9]|1[0-2])", n)
    if m:
        return MONTH_ALIASES[m.group(1).zfill(2)]
    return None


def detect_month_columns(df: pd.DataFrame) -> Dict[str, str]:
    found = {}
    for col in df.columns:
        m = find_month_in_col(col)
        if m and m not in found:
            nc = norm_txt(col)
            if any(bad in nc for bad in ["ytd", "bytd", "fy", "total", "var"]):
                continue
            found[m] = col
    return {m: found[m] for m in MONTH_ORDER if m in found}


def find_best_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    normalized = {norm_txt(c): c for c in df.columns}
    for cand in candidates:
        if norm_txt(cand) in normalized:
            return normalized[norm_txt(cand)]
    for col in df.columns:
        nc = norm_txt(col)
        if any(norm_txt(cand) in nc for cand in candidates):
            return col
    return None


def money(value: float) -> str:
    try:
        value = float(value)
    except Exception:
        value = 0.0
    if abs(value) >= 1_000_000:
        return f"US$ {value/1_000_000:,.1f} MM"
    if abs(value) >= 1_000:
        return f"US$ {value/1_000:,.1f} k"
    return f"US$ {value:,.0f}"


def safe_div(num: float, den: float, default: float = 1.0) -> float:
    if pd.isna(den) or abs(float(den)) < 1e-9:
        return default
    return float(num) / float(den)


def scenario_curve(n: int, intensity: float, steepness: float) -> np.ndarray:
    if n <= 0:
        return np.array([])
    x = np.linspace(-2, 2, n)
    sig = 1 / (1 + np.exp(-steepness * x))
    centered = sig - sig.mean()
    return 1 + intensity * centered


def context_key(value: object) -> str:
    s = norm_txt(value)
    if "fuel" in s or "combust" in s or "diesel" in s:
        return "Fuel"
    if "contract" in s or "contrat" in s or "tercer" in s:
        return "Contractors"
    if "maint" in s or "mant" in s:
        return "Maintenance"
    if "labor" in s or "remun" in s or "sueldo" in s or "salary" in s:
        return "Labor"
    if "power" in s or "energia" in s or "electric" in s:
        return "Power"
    if "spare" in s or "repuesto" in s:
        return "Spare Parts"
    if "rehandling" in s or "remanejo" in s:
        return "Rehandling"
    return "Other"


def clean_sheet_headers(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    cols_norm = [norm_txt(c) for c in df.columns]
    if "resp" in cols_norm and any("jan" in c or "ene" in c for c in cols_norm):
        return df
    max_rows = min(10, len(df))
    for i in range(max_rows):
        row_values = [norm_txt(x) for x in df.iloc[i].tolist()]
        has_resp = "resp" in row_values
        has_cc = "cc" in row_values
        has_month = any(
            any(alias in cell for alias in ["jan", "ene", "feb", "mar", "apr", "abr", "may", "jun", "jul", "aug", "ago", "sep", "oct", "nov", "dec", "dic"])
            for cell in row_values
        )
        if has_resp and has_cc and has_month:
            new_cols = df.iloc[i].astype(str).tolist()
            df = df.iloc[i + 1:].copy()
            df.columns = new_cols
            df = df.dropna(how="all")
            df = df.loc[:, ~pd.Index(df.columns).astype(str).str.startswith("nan")]
            return df.reset_index(drop=True)
    return df


# ─────────────────────────────────────────────
# CARGA DE DATOS
# ─────────────────────────────────────────────

@st.cache_data(show_spinner=False)
def read_excel_sheets(uploaded_file) -> Dict[str, pd.DataFrame]:
    xls = pd.ExcelFile(uploaded_file)
    sheets = {}
    for sheet in xls.sheet_names:
        raw = pd.read_excel(uploaded_file, sheet_name=sheet)
        sheets[sheet] = clean_sheet_headers(raw)
    return sheets


# ─────────────────────────────────────────────
# PREPARACIÓN Y FORECAST
# ─────────────────────────────────────────────

def prepare_data(
    gastos_raw: pd.DataFrame,
    budget_raw: Optional[pd.DataFrame],
    gastos_month_cols: Dict[str, str],
    budget_month_cols: Dict[str, str],
    actual_months: List[str],
    forecast_months: List[str],
    classif_col: Optional[str],
    cc_col_gastos: Optional[str],
    cc_col_budget: Optional[str],
    budget_fy_col: Optional[str],
    extra_dims: List[str],
) -> pd.DataFrame:
    df = gastos_raw.copy()
    for month, col in gastos_month_cols.items():
        df[f"Actual_{month}"] = pd.to_numeric(df[col], errors="coerce").fillna(0)

    if classif_col and classif_col in df.columns:
        df["Naturaleza"] = df[classif_col].fillna("Other").astype(str)
    else:
        df["Naturaleza"] = "Other"
    df["Contexto_Mina"] = df["Naturaleza"].map(context_key)

    if cc_col_gastos and cc_col_gastos in df.columns:
        df["_CC_KEY_"] = df[cc_col_gastos].astype(str).str.replace(".0", "", regex=False).str.strip()
    else:
        df["_CC_KEY_"] = np.arange(len(df)).astype(str)

    for m in MONTH_ORDER:
        df[f"Budget_{m}"] = 0.0

    if budget_raw is not None and budget_month_cols:
        b = budget_raw.copy()
        if cc_col_budget and cc_col_budget in b.columns:
            b["_CC_KEY_"] = b[cc_col_budget].astype(str).str.replace(".0", "", regex=False).str.strip()
        else:
            b["_CC_KEY_"] = np.arange(len(b)).astype(str)
        agg_dict = {}
        for m, col in budget_month_cols.items():
            b[f"Budget_{m}"] = pd.to_numeric(b[col], errors="coerce").fillna(0)
            agg_dict[f"Budget_{m}"] = "sum"
        if budget_fy_col and budget_fy_col in b.columns:
            b["Budget_FY_Input"] = pd.to_numeric(b[budget_fy_col], errors="coerce").fillna(0)
            agg_dict["Budget_FY_Input"] = "sum"
        b_agg = b.groupby("_CC_KEY_", as_index=False).agg(agg_dict)
        df = df.merge(b_agg, on="_CC_KEY_", how="left", suffixes=("", "_from_budget"))
        for m in MONTH_ORDER:
            merged = f"Budget_{m}_from_budget"
            if merged in df.columns:
                df[f"Budget_{m}"] = pd.to_numeric(df[merged], errors="coerce").fillna(df[f"Budget_{m}"])
                df.drop(columns=[merged], inplace=True)
    else:
        for col in gastos_raw.columns:
            nc = norm_txt(col)
            if "budget" in nc or "presupuesto" in nc:
                m = find_month_in_col(col)
                if m:
                    df[f"Budget_{m}"] = pd.to_numeric(gastos_raw[col], errors="coerce").fillna(0)

    for m in MONTH_ORDER:
        if f"Budget_{m}" not in df.columns:
            df[f"Budget_{m}"] = 0.0
        df[f"Budget_{m}"] = pd.to_numeric(df[f"Budget_{m}"], errors="coerce").fillna(0)

    df["Actual_YTD"] = df[[f"Actual_{m}" for m in actual_months if f"Actual_{m}" in df.columns]].sum(axis=1)
    df["Budget_YTD"] = df[[f"Budget_{m}" for m in actual_months]].sum(axis=1)
    df["Budget_Remaining"] = df[[f"Budget_{m}" for m in forecast_months]].sum(axis=1)
    df["Budget_FY_Model"] = df[[f"Budget_{m}" for m in MONTH_ORDER]].sum(axis=1)
    if "Budget_FY_Input" in df.columns:
        df["Budget_FY_Model"] = df["Budget_FY_Input"].where(df["Budget_FY_Input"].abs() > 1e-9, df["Budget_FY_Model"])

    df["Budget_YTD"] = np.where(df["Budget_YTD"].abs() < 1e-9, df["Actual_YTD"], df["Budget_YTD"])
    return df


def calculate_forecast(
    df: pd.DataFrame,
    actual_months: List[str],
    forecast_months: List[str],
    sensitivity_mult: float,
    momentum_mult: float,
    scenario_mult: float,
    steepness: float,
) -> pd.DataFrame:
    out = df.copy()
    if len(actual_months) == 0 or len(forecast_months) == 0:
        raise ValueError("Debes tener al menos 1 mes real y 1 mes forecast.")

    recent = actual_months[-min(2, len(actual_months)):]
    early = actual_months[:min(3, len(actual_months))]
    out["Prom_Reciente"] = out[[f"Actual_{m}" for m in recent]].mean(axis=1)
    out["Prom_Inicial"] = out[[f"Actual_{m}" for m in early]].mean(axis=1)
    out["Factor_Ejecucion"] = [safe_div(a, b, 1.0) for a, b in zip(out["Actual_YTD"], out["Budget_YTD"])]
    out["Factor_Tendencia_Bruto"] = [safe_div(a, b, 1.0) for a, b in zip(out["Prom_Reciente"], out["Prom_Inicial"])]
    out["Factor_Ejecucion"] = out["Factor_Ejecucion"].clip(0.2, 2.8)
    out["Factor_Tendencia_Bruto"] = out["Factor_Tendencia_Bruto"].clip(0.35, 2.5)

    for m in forecast_months:
        out[f"Forecast_{m}"] = 0.0

    for idx, row in out.iterrows():
        params = DEFAULT_CONTEXT.get(row["Contexto_Mina"], DEFAULT_CONTEXT["Other"])
        sens = params["sens"] * sensitivity_mult
        mom = params["mom"] * momentum_mult
        curve_intensity = params["curve"] * scenario_mult
        min_f, max_f = params["min"], params["max"]
        base_factor = 1 + sens * (row["Factor_Ejecucion"] - 1) + mom * (row["Factor_Tendencia_Bruto"] - 1)
        base_factor = float(np.clip(base_factor, min_f, max_f))
        curve = scenario_curve(len(forecast_months), curve_intensity, steepness)
        for j, m in enumerate(forecast_months):
            budget = row.get(f"Budget_{m}", 0.0)
            fallback = row["Actual_YTD"] / max(len(actual_months), 1)
            base = budget if abs(budget) > 1e-9 else fallback
            out.at[idx, f"Forecast_{m}"] = base * base_factor * curve[j]

    out["Forecast_Remaining"] = out[[f"Forecast_{m}" for m in forecast_months]].sum(axis=1)
    out["Forecast_FY_Modelo"] = out["Actual_YTD"] + out["Forecast_Remaining"]
    out["Var_vs_Budget"] = out["Forecast_FY_Modelo"] - out["Budget_FY_Model"]
    out["Var_vs_Budget_%"] = [safe_div(v, b, 0.0) for v, b in zip(out["Var_vs_Budget"], out["Budget_FY_Model"])]
    out["Recomendacion"] = np.select(
        [out["Var_vs_Budget_%"] > 0.10, out["Var_vs_Budget_%"] < -0.10],
        ["Riesgo de sobreconsumo: activar plan de control y revisión contractual.", "Subejecución relevante: revisar reprogramación operacional o liberación de presupuesto."],
        default="Desviación controlada: mantener monitoreo mensual.",
    )
    out["Justificacion_Mina"] = out["Contexto_Mina"].map(lambda k: DEFAULT_CONTEXT.get(k, DEFAULT_CONTEXT["Other"])["why"])
    return out


def aggregate_for_charts(df: pd.DataFrame, group_col: str) -> pd.DataFrame:
    if group_col not in df.columns:
        group_col = "Naturaleza"
    return df.groupby(group_col, dropna=False, as_index=False).agg(
        Actual_YTD=("Actual_YTD", "sum"),
        Budget_FY_Model=("Budget_FY_Model", "sum"),
        Forecast_FY_Modelo=("Forecast_FY_Modelo", "sum"),
        Var_vs_Budget=("Var_vs_Budget", "sum"),
    ).sort_values("Forecast_FY_Modelo", ascending=False)


# ─────────────────────────────────────────────
# SIMULACIÓN 5 AÑOS
# ─────────────────────────────────────────────

def calcular_estacionalidad_mensual(
    df: pd.DataFrame,
    actual_months: List[str],
    forecast_months: List[str],
) -> Dict[str, float]:
    """Calcula el % que representa cada mes sobre el total FY (Actual + Forecast).
    Se usa como patrón de estacionalidad para repartir Año_1 en 12 meses."""
    total_fy = 0.0
    montos_mes = {}
    for m in MONTH_ORDER:
        if m in actual_months:
            col = f"Actual_{m}"
        elif m in forecast_months:
            col = f"Forecast_{m}"
        else:
            col = None
        monto = float(df[col].sum()) if col and col in df.columns else 0.0
        montos_mes[m] = monto
        total_fy += monto
    if total_fy <= 1e-9:
        return {m: 1 / 12 for m in MONTH_ORDER}
    return {m: montos_mes[m] / total_fy for m in MONTH_ORDER}


def simular_5_anos(df_base: pd.DataFrame, inf_anual: float, crec_ops: float) -> pd.DataFrame:
    """Proyecta 5 años usando el Forecast actual como base y tasas de crecimiento compuestas."""
    df_5y = df_base.copy()

    def obtener_tasa(ctx):
        if ctx in ["Labor", "Other"]:
            return inf_anual / 100
        if ctx in ["Fuel", "Power", "Spare Parts", "Rehandling"]:
            return (inf_anual + crec_ops) / 100
        return (inf_anual + (crec_ops * 0.5)) / 100

    df_5y["Tasa_Crecimiento"] = df_5y["Contexto_Mina"].apply(obtener_tasa)
    df_5y["Año_0_Base"] = df_5y["Forecast_FY_Modelo"]

    # Fix: prev_col evita buscar "Año_0" que no existe
    prev_col = "Año_0_Base"
    for i in range(1, 6):
        new_col = f"Año_{i}"
        df_5y[new_col] = df_5y[prev_col] * (1 + df_5y["Tasa_Crecimiento"])
        prev_col = new_col

    return df_5y


# ─────────────────────────────────────────────
# EXPORTACIÓN EXCEL
# ─────────────────────────────────────────────

def to_excel_bytes(df: pd.DataFrame, summary: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Forecast_Detalle")
        summary.to_excel(writer, index=False, sheet_name="Resumen")
        wb = writer.book
        money_fmt = wb.add_format({"num_format": '#,##0;[Red]-#,##0'})
        pct_fmt = wb.add_format({"num_format": '0.0%'})
        header_fmt = wb.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1})
        for sheet_name in ["Forecast_Detalle", "Resumen"]:
            ws = writer.sheets[sheet_name]
            for c, col in enumerate((df if sheet_name == "Forecast_Detalle" else summary).columns):
                ws.write(0, c, col, header_fmt)
                width = min(max(len(str(col)) + 2, 12), 35)
                ws.set_column(c, c, width)
                if any(k in str(col).lower() for k in ["budget", "forecast", "actual", "var", "prom"]):
                    ws.set_column(c, c, width, money_fmt)
                if "%" in str(col):
                    ws.set_column(c, c, width, pct_fmt)
    return output.getvalue()


def to_excel_completo(
    df_forecast: pd.DataFrame,
    summary_forecast: pd.DataFrame,
    resumen_5y: pd.DataFrame,
    resumen_largo: pd.DataFrame,
    actual_months: List[str],
    forecast_months: List[str],
) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        wb = writer.book
        money_fmt  = wb.add_format({"num_format": '#,##0;[Red]-#,##0'})
        pct_fmt    = wb.add_format({"num_format": '0.0%'})
        header_fmt = wb.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1})
        green_fmt  = wb.add_format({"bold": True, "bg_color": "#C6EFCE", "border": 1})
        red_fmt    = wb.add_format({"bold": True, "bg_color": "#FFC7CE", "border": 1})

        def write_sheet(df, sheet_name, money_keys=(), pct_keys=()):
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            ws = writer.sheets[sheet_name]
            for c, col in enumerate(df.columns):
                ws.write(0, c, col, header_fmt)
                width = min(max(len(str(col)) + 2, 14), 38)
                if any(k in str(col).lower() for k in pct_keys):
                    ws.set_column(c, c, width, pct_fmt)
                elif any(k in str(col).lower() for k in money_keys):
                    ws.set_column(c, c, width, money_fmt)
                else:
                    ws.set_column(c, c, width)

        # Hoja 1: Resumen Forecast
        write_sheet(
            summary_forecast, "Resumen_Forecast",
            money_keys=["actual", "budget", "forecast", "var"],
            pct_keys=["%"],
        )

        # Hoja 2: Detalle Forecast fila a fila
        cols_detalle = (
            ["Naturaleza", "Contexto_Mina"]
            + [f"Actual_{m}"   for m in actual_months   if f"Actual_{m}"   in df_forecast.columns]
            + [f"Budget_{m}"   for m in actual_months   if f"Budget_{m}"   in df_forecast.columns]
            + [f"Forecast_{m}" for m in forecast_months if f"Forecast_{m}" in df_forecast.columns]
            + ["Actual_YTD", "Budget_FY_Model", "Forecast_FY_Modelo",
               "Var_vs_Budget", "Var_vs_Budget_%", "Recomendacion"]
        )
        cols_detalle = list(dict.fromkeys([c for c in cols_detalle if c in df_forecast.columns]))
        write_sheet(
            df_forecast[cols_detalle], "Detalle_Forecast",
            money_keys=["actual", "budget", "forecast", "var", "prom"],
            pct_keys=["%"],
        )

        # Colorear Var_vs_Budget_% en Detalle_Forecast
        ws2 = writer.sheets["Detalle_Forecast"]
        if "Var_vs_Budget_%" in cols_detalle:
            col_idx = cols_detalle.index("Var_vs_Budget_%")
            for row_idx, val in enumerate(df_forecast["Var_vs_Budget_%"], start=1):
                try:
                    v = float(val)
                    fmt = red_fmt if v > 0.10 else (green_fmt if v < -0.10 else pct_fmt)
                except Exception:
                    fmt = pct_fmt
                ws2.write(row_idx, col_idx, val, fmt)

        # Hoja 3: Simulación LRP 5 Años (matriz con meses de Año_1)
        write_sheet(
            resumen_5y, "Simulacion_5_Anos",
            money_keys=["fy", "jan", "feb", "mar", "apr", "may", "jun",
                        "jul", "aug", "sep", "oct", "nov", "dec"],
        )

        # Hoja 4: LRP formato largo (para pivots)
        write_sheet(
            resumen_largo, "LRP_Detalle_Largo",
            money_keys=["presupuesto"],
        )

        # Hoja 5: Parámetros del modelo
        params_data = {
            "Parámetro": [
                "Último mes real (YTD)", "Meses reales", "Meses forecast",
                "Sensibilidad ejecución YTD", "Peso tendencia reciente",
                "Intensidad curva no lineal", "Forma curva logística",
                "Inflación anual estimada (%)", "Crecimiento operacional anual (%)",
            ],
            "Valor": [
                cutoff_month, ", ".join(actual_months), ", ".join(forecast_months),
                sensitivity_mult, momentum_mult, scenario_mult, steepness,
                cagr_inf, cagr_ops,
            ],
        }
        pd.DataFrame(params_data).to_excel(writer, index=False, sheet_name="Parametros_Modelo")
        ws5 = writer.sheets["Parametros_Modelo"]
        ws5.set_column(0, 0, 35)
        ws5.set_column(1, 1, 45)
        for c, col in enumerate(["Parámetro", "Valor"]):
            ws5.write(0, c, col, header_fmt)

    return output.getvalue()


# ═══════════════════════════════════════════════════════════════
# INTERFAZ STREAMLIT
# ═══════════════════════════════════════════════════════════════

st.title("⛏️ Forecast 5+7 No Lineal Dinámico")
st.caption("Aplicación web para presupuesto y gastos mineros: carga Excel, detecta meses, proyecta forecast y genera hallazgos ejecutivos.")

with st.expander("📌 Metodología del modelo", expanded=False):
    st.markdown(
        """
        **Forecast mensual futuro = Budget mensual × Factor de ejecución × Factor de tendencia × Curva no lineal**

        - **Factor de ejecución:** compara gasto real acumulado contra presupuesto acumulado.
        - **Factor de tendencia:** compara meses recientes contra meses iniciales.
        - **Curva no lineal:** distribuye el ajuste en los meses restantes con una forma logística, evitando una simple extrapolación lineal.
        - **Contexto minero:** los factores cambian según naturaleza del gasto: Fuel, Contractors, Maintenance, Labor, Power, Spare Parts u Other.
        """
    )

uploaded = st.sidebar.file_uploader("Sube tu Excel de gastos/presupuesto", type=["xlsx", "xls"])

if uploaded is None:
    st.info("Sube un archivo Excel para comenzar. La app funciona con hojas similares a 'Gastos' y 'Budget', pero permite seleccionar otra estructura.")
    st.stop()

try:
    sheets = read_excel_sheets(uploaded)
except Exception as e:
    st.error(f"No pude leer el Excel: {e}")
    st.stop()

sheet_names = list(sheets.keys())
def_sheet_gastos = next((s for s in sheet_names if norm_txt(s) in ["gastos", "expenses", "costs"]), sheet_names[0])
def_sheet_budget = next((s for s in sheet_names if "budget" in norm_txt(s) or "presupuesto" in norm_txt(s)), sheet_names[0])

st.sidebar.subheader("1) Selección de hojas")
gastos_sheet = st.sidebar.selectbox("Hoja de gastos reales", sheet_names, index=sheet_names.index(def_sheet_gastos))
budget_sheet = st.sidebar.selectbox("Hoja de presupuesto", ["Sin hoja Budget / usar misma hoja"] + sheet_names, index=(sheet_names.index(def_sheet_budget) + 1 if def_sheet_budget in sheet_names else 0))

gastos_raw = sheets[gastos_sheet].copy()
budget_raw = None if budget_sheet.startswith("Sin hoja") else sheets[budget_sheet].copy()

gastos_month_cols = detect_month_columns(gastos_raw)
budget_month_cols = detect_month_columns(budget_raw) if budget_raw is not None else {}

if len(gastos_month_cols) < 2:
    st.error("No detecté suficientes columnas mensuales en la hoja de gastos. Revisa que existan columnas tipo Jan-25, Ene-25, 2025-01, etc.")
    st.write("Columnas detectadas:", list(gastos_raw.columns))
    st.stop()

available_months = [m for m in MONTH_ORDER if m in gastos_month_cols]

st.sidebar.subheader("2) Configuración Forecast")
cutoff_default = available_months[min(4, len(available_months)-2)] if len(available_months) >= 6 else available_months[-2]
cutoff_month = st.sidebar.selectbox("Último mes real / cierre YTD", available_months[:-1], index=available_months[:-1].index(cutoff_default) if cutoff_default in available_months[:-1] else 0)
cutoff_idx = MONTH_ORDER.index(cutoff_month)
actual_months = [m for m in MONTH_ORDER[:cutoff_idx+1] if m in gastos_month_cols]
forecast_months = [m for m in MONTH_ORDER[cutoff_idx+1:] if m in MONTH_ORDER]

st.sidebar.write(f"Meses reales: {', '.join(actual_months)}")
st.sidebar.write(f"Meses forecast: {', '.join(forecast_months)}")

st.sidebar.subheader("3) Columnas de negocio")
all_cols = [None] + list(gastos_raw.columns)
classif_guess = find_best_col(gastos_raw, POSSIBLE_CLASSIF)
cc_guess_g = find_best_col(gastos_raw, POSSIBLE_CC)
cc_guess_b = find_best_col(budget_raw, POSSIBLE_CC) if budget_raw is not None else None
budget_fy_guess = find_best_col(budget_raw, POSSIBLE_BUDGET_FY) if budget_raw is not None else None

classif_col = st.sidebar.selectbox("Columna de naturaleza/clasificación", all_cols, index=all_cols.index(classif_guess) if classif_guess in all_cols else 0)
cc_col_gastos = st.sidebar.selectbox("Llave en gastos (CC/Resp/Centro costo)", all_cols, index=all_cols.index(cc_guess_g) if cc_guess_g in all_cols else 0)
if budget_raw is not None:
    budget_cols_options = [None] + list(budget_raw.columns)
    cc_col_budget = st.sidebar.selectbox("Llave en Budget", budget_cols_options, index=budget_cols_options.index(cc_guess_b) if cc_guess_b in budget_cols_options else 0)
    budget_fy_col = st.sidebar.selectbox("Columna Budget FY total (opcional)", budget_cols_options, index=budget_cols_options.index(budget_fy_guess) if budget_fy_guess in budget_cols_options else 0)
else:
    cc_col_budget, budget_fy_col = None, None

candidate_dims = [c for c in gastos_raw.columns if c not in gastos_month_cols.values()]
default_dims = [c for c in candidate_dims if norm_txt(c) in ["vp", "gerencia", "desc resp", "desc proc", "classif", "naturaleza"]]
extra_dims = st.sidebar.multiselect("Filtros/dimensiones para dashboard", candidate_dims, default=default_dims[:6])

st.sidebar.subheader("4) Parámetros del modelo")
sensitivity_mult = st.sidebar.slider("Sensibilidad a ejecución YTD", 0.2, 2.0, 1.0, 0.05)
momentum_mult = st.sidebar.slider("Peso de tendencia reciente", 0.0, 2.0, 1.0, 0.05)
scenario_mult = st.sidebar.slider("Intensidad curva no lineal", 0.0, 2.0, 1.0, 0.05)
steepness = st.sidebar.slider("Forma de curva logística", 0.5, 3.0, 1.35, 0.05)

st.sidebar.subheader("5) Simulación Estratégica (5 Años)")
st.sidebar.markdown("Define el escenario macroeconómico y operativo:")
cagr_inf = st.sidebar.slider("Inflación Anual Estimada (%)", 0.0, 10.0, 3.0, 0.5)
cagr_ops = st.sidebar.slider("Crecimiento Operacional Anual (%)", -5.0, 15.0, 2.0, 0.5)

try:
    prepared = prepare_data(
        gastos_raw, budget_raw, gastos_month_cols, budget_month_cols,
        actual_months, forecast_months, classif_col, cc_col_gastos, cc_col_budget,
        budget_fy_col, extra_dims,
    )
    model = calculate_forecast(prepared, actual_months, forecast_months, sensitivity_mult, momentum_mult, scenario_mult, steepness)
except Exception as e:
    st.error(f"Error al preparar/calcular el forecast: {e}")
    st.stop()

# Filtros dinámicos
filtered = model.copy()
with st.sidebar.expander("6) Filtros", expanded=True):
    if "Naturaleza" in filtered.columns:
        selected_nat = st.multiselect("Naturaleza", sorted(filtered["Naturaleza"].dropna().astype(str).unique()), default=[])
        if selected_nat:
            filtered = filtered[filtered["Naturaleza"].astype(str).isin(selected_nat)]
    for dim in extra_dims[:8]:
        if dim in filtered.columns:
            vals = sorted(filtered[dim].dropna().astype(str).unique())
            if 1 < len(vals) <= 200:
                sel = st.multiselect(str(dim), vals, default=[])
                if sel:
                    filtered = filtered[filtered[dim].astype(str).isin(sel)]

# ─────────────────────────────────────────────
# KPIs
# ─────────────────────────────────────────────
kpi1, kpi2, kpi3, kpi4 = st.columns(4)
actual_total = filtered["Actual_YTD"].sum()
budget_total = filtered["Budget_FY_Model"].sum()
forecast_total = filtered["Forecast_FY_Modelo"].sum()
var_total = forecast_total - budget_total
var_pct = safe_div(var_total, budget_total, 0)
kpi1.metric("Actual YTD", money(actual_total))
kpi2.metric("Budget FY", money(budget_total))
kpi3.metric("Forecast FY Modelo", money(forecast_total))
kpi4.metric("Var vs Budget", money(var_total), f"{var_pct:.1%}")

st.divider()

# Agrupación para gráficos
chart_dim_options = ["Naturaleza", "Contexto_Mina"] + [c for c in extra_dims if c in filtered.columns]
chart_dim = st.selectbox("Agrupar dashboard por", chart_dim_options, index=0)
summary = aggregate_for_charts(filtered, chart_dim)

# ─────────────────────────────────────────────
# GENERACIÓN DE INFORME PDF (ES + EN)
# ─────────────────────────────────────────────

def generar_pdf_informe(
    lang: str,
    actual_total: float,
    budget_total: float,
    forecast_total: float,
    var_total: float,
    var_pct: float,
    summary: pd.DataFrame,
    chart_dim: str,
    actual_months: List[str],
    forecast_months: List[str],
    cutoff_month: str,
    sensitivity_mult: float,
    momentum_mult: float,
    cagr_inf: float,
    cagr_ops: float,
    resumen_5y: pd.DataFrame,
    registro: List[dict],
) -> bytes:
    """Genera informe ejecutivo en PDF en español o inglés."""

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.8*cm, leftMargin=1.8*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    # ── Estilos ────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    COLOR_HEADER = colors.HexColor("#1F4E79")
    COLOR_ACCENT = colors.HexColor("#2E75B6")
    COLOR_LIGHT  = colors.HexColor("#D9EAF7")
    COLOR_GREEN  = colors.HexColor("#C6EFCE")
    COLOR_RED    = colors.HexColor("#FFC7CE")
    COLOR_GRAY   = colors.HexColor("#F2F2F2")

    s_title = ParagraphStyle("s_title", parent=styles["Title"],
        fontSize=20, textColor=COLOR_HEADER, spaceAfter=6, alignment=TA_CENTER, fontName="Helvetica-Bold")
    s_subtitle = ParagraphStyle("s_subtitle", parent=styles["Normal"],
        fontSize=11, textColor=COLOR_ACCENT, spaceAfter=14, alignment=TA_CENTER, fontName="Helvetica")
    s_h1 = ParagraphStyle("s_h1", parent=styles["Heading1"],
        fontSize=13, textColor=COLOR_HEADER, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold",
        borderPad=4, borderColor=COLOR_ACCENT, borderWidth=0)
    s_h2 = ParagraphStyle("s_h2", parent=styles["Heading2"],
        fontSize=11, textColor=COLOR_ACCENT, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    s_body = ParagraphStyle("s_body", parent=styles["Normal"],
        fontSize=9, spaceAfter=5, leading=13, fontName="Helvetica")
    s_body_bold = ParagraphStyle("s_body_bold", parent=s_body, fontName="Helvetica-Bold")
    s_caption = ParagraphStyle("s_caption", parent=styles["Normal"],
        fontSize=8, textColor=colors.gray, spaceAfter=4, fontName="Helvetica")
    s_kpi_label = ParagraphStyle("s_kpi_label", parent=styles["Normal"],
        fontSize=8, textColor=colors.white, alignment=TA_CENTER, fontName="Helvetica")
    s_kpi_value = ParagraphStyle("s_kpi_value", parent=styles["Normal"],
        fontSize=14, textColor=colors.white, alignment=TA_CENTER, fontName="Helvetica-Bold")

    # ── Textos bilingues ───────────────────────────────────────
    T = {
        "es": {
            "title": "Informe Ejecutivo — Forecast 5+7 No Lineal Dinámico",
            "subtitle": "Análisis de Presupuesto y Proyección Minera",
            "date_label": "Fecha de generación:",
            "period": f"Período real: Jan — {cutoff_month} | Período forecast: {forecast_months[0] if forecast_months else '—'} — {forecast_months[-1] if forecast_months else '—'}",
            "sec_kpi": "1. KPIs Ejecutivos",
            "kpi_actual": "Actual YTD",
            "kpi_budget": "Budget FY",
            "kpi_forecast": "Forecast FY Modelo",
            "kpi_var": "Variación vs Budget",
            "sec_model": "2. Parámetros del Modelo",
            "param_cutoff": "Último mes real (YTD)",
            "param_actual": "Meses reales",
            "param_forecast": "Meses forecast",
            "param_sens": "Sensibilidad ejecución YTD",
            "param_mom": "Peso tendencia reciente",
            "param_inf": "Inflación anual estimada",
            "param_ops": "Crecimiento operacional anual",
            "sec_summary": "3. Resumen por Categoría",
            "col_cat": "Categoría",
            "col_actual": "Actual YTD",
            "col_budget": "Budget FY",
            "col_forecast": "Forecast FY",
            "col_var": "Var vs Budget",
            "col_var_pct": "Var %",
            "sec_lrp": "4. Simulación LRP 5 Años (US$ MM)",
            "sec_hallazgos": "5. Hallazgos y Recomendaciones",
            "sec_sens": "6. Registro de Sensibilidades",
            "no_sens": "No se registraron escenarios de sensibilidad en esta sesión.",
            "methodology": "Metodología del Modelo",
            "method_body": (
                "El modelo Forecast 5+7 proyecta el cierre anual combinando: (1) Factor de ejecución: "
                "gasto real acumulado vs presupuesto acumulado. (2) Factor de tendencia: meses recientes "
                "vs meses iniciales. (3) Curva no lineal logística distribuida en los meses restantes. "
                "(4) Contexto minero: parámetros diferenciados según naturaleza del gasto (Fuel, Labor, "
                "Maintenance, Power, Contractors, Spare Parts, Rehandling, Other)."
            ),
            "footer": "Modelo Forecast 5+7 No Lineal Dinámico — Uso interno de gestión presupuestaria",
            "over": "SOBREEJECUCIÓN",
            "under": "SUBEJECUCIÓN",
            "on_track": "ALINEADO",
            "hallazgo_1": f"El Forecast FY Modelo proyecta un cierre de {money(forecast_total)}, representando una desviación de {var_pct:+.1%} respecto al Budget FY de {money(budget_total)}.",
            "hallazgo_2": "Las categorías con mayor peso presupuestario concentran las principales desviaciones y deben priorizarse en el control operacional.",
            "hallazgo_3": f"La simulación LRP a 5 años incorpora una inflación de {cagr_inf}% y crecimiento operacional de {cagr_ops}%, proyectando la estructura de costos hasta el año {datetime.datetime.now().year + 5}.",
            "rec_1": "Activar plan de control en categorías con desviación positiva superior al 10%.",
            "rec_2": "Revisar reprogramación operacional en partidas con subejecución relevante.",
            "rec_3": "Implementar alertas tempranas mensuales por naturaleza de gasto y centro de costo.",
        },
        "en": {
            "title": "Executive Report — Dynamic Non-Linear 5+7 Forecast",
            "subtitle": "Mining Budget Analysis & Projection",
            "date_label": "Generated on:",
            "period": f"Actual period: Jan — {cutoff_month} | Forecast period: {forecast_months[0] if forecast_months else '—'} — {forecast_months[-1] if forecast_months else '—'}",
            "sec_kpi": "1. Executive KPIs",
            "kpi_actual": "Actual YTD",
            "kpi_budget": "Budget FY",
            "kpi_forecast": "Forecast FY Model",
            "kpi_var": "Variance vs Budget",
            "sec_model": "2. Model Parameters",
            "param_cutoff": "Last actual month (YTD)",
            "param_actual": "Actual months",
            "param_forecast": "Forecast months",
            "param_sens": "YTD execution sensitivity",
            "param_mom": "Recent trend weight",
            "param_inf": "Estimated annual inflation",
            "param_ops": "Annual operational growth",
            "sec_summary": "3. Summary by Category",
            "col_cat": "Category",
            "col_actual": "Actual YTD",
            "col_budget": "Budget FY",
            "col_forecast": "Forecast FY",
            "col_var": "Var vs Budget",
            "col_var_pct": "Var %",
            "sec_lrp": "4. 5-Year LRP Simulation (US$ MM)",
            "sec_hallazgos": "5. Findings & Recommendations",
            "sec_sens": "6. Sensitivity Scenarios Log",
            "no_sens": "No sensitivity scenarios were recorded in this session.",
            "methodology": "Model Methodology",
            "method_body": (
                "The 5+7 Forecast model projects the annual close by combining: (1) Execution factor: "
                "actual YTD spend vs YTD budget. (2) Trend factor: recent months vs early months. "
                "(3) Non-linear logistic curve distributed over remaining months. "
                "(4) Mining context: differentiated parameters by cost nature (Fuel, Labor, "
                "Maintenance, Power, Contractors, Spare Parts, Rehandling, Other)."
            ),
            "footer": "Dynamic Non-Linear 5+7 Forecast Model — Internal budget management use",
            "over": "OVER BUDGET",
            "under": "UNDER BUDGET",
            "on_track": "ON TRACK",
            "hallazgo_1": f"The FY Model Forecast projects a year-end close of {money(forecast_total)}, representing a {var_pct:+.1%} deviation vs the FY Budget of {money(budget_total)}.",
            "hallazgo_2": "Categories with the highest budget weight concentrate the main deviations and should be prioritized in operational control.",
            "hallazgo_3": f"The 5-year LRP simulation incorporates {cagr_inf}% inflation and {cagr_ops}% operational growth, projecting the cost structure through year {datetime.datetime.now().year + 5}.",
            "rec_1": "Activate control plan for categories with positive deviation above 10%.",
            "rec_2": "Review operational rescheduling for categories with significant under-execution.",
            "rec_3": "Implement monthly early-warning alerts by cost nature and cost center.",
        },
    }
    t = T[lang]

    # ── Helpers de tabla ───────────────────────────────────────
    def make_table(data, col_widths, header_bg=COLOR_LIGHT):
        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        style = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR",  (0, 0), (-1, 0), COLOR_HEADER),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, 0), 8),
            ("ALIGN",      (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME",   (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",   (0, 1), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, COLOR_GRAY]),
            ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ])
        tbl.setStyle(style)
        return tbl

    def kpi_box(label, value, color=COLOR_ACCENT):
        return Table(
            [[Paragraph(label, s_kpi_label)], [Paragraph(value, s_kpi_value)]],
            colWidths=[4.5*cm],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), color),
                ("ROWBACKGROUNDS", (0, 0), (0, -1), [color, color]),
                ("ALIGN",   (0, 0), (0, -1), "CENTER"),
                ("VALIGN",  (0, 0), (0, -1), "MIDDLE"),
                ("TOPPADDING",    (0, 0), (0, -1), 8),
                ("BOTTOMPADDING", (0, 0), (0, -1), 8),
                ("LEFTPADDING",   (0, 0), (0, -1), 4),
                ("RIGHTPADDING",  (0, 0), (0, -1), 4),
                ("ROUNDEDCORNERS", [4]),
            ])
        )

    # ── Construcción del documento ─────────────────────────────
    story = []

    # Portada
    story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph("⛏️", ParagraphStyle("icon", fontSize=32, alignment=TA_CENTER, spaceAfter=8)))
    story.append(Paragraph(t["title"], s_title))
    story.append(Paragraph(t["subtitle"], s_subtitle))
    story.append(HRFlowable(width="100%", thickness=2, color=COLOR_ACCENT, spaceAfter=8))
    story.append(Paragraph(f"{t['date_label']} {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", s_caption))
    story.append(Paragraph(t["period"], s_caption))
    story.append(Spacer(1, 0.5*cm))

    # ── Sección 1: KPIs ────────────────────────────────────────
    story.append(Paragraph(t["sec_kpi"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))

    kpi_color = COLOR_RED if var_total > 0 else (COLOR_GREEN if var_total < 0 else COLOR_ACCENT)
    kpi_color_rl = colors.HexColor("#C00000") if var_total > 0 else (colors.HexColor("#375623") if var_total < 0 else COLOR_ACCENT)
    var_label = t["over"] if var_total > 0 else (t["under"] if var_total < 0 else t["on_track"])

    kpi_row = Table(
        [[
            kpi_box(t["kpi_actual"],   money(actual_total),   COLOR_ACCENT),
            kpi_box(t["kpi_budget"],   money(budget_total),   COLOR_ACCENT),
            kpi_box(t["kpi_forecast"], money(forecast_total), COLOR_ACCENT),
            kpi_box(f"{t['kpi_var']}\n{var_label}", f"{var_pct:+.1%}\n{money(var_total)}", kpi_color_rl),
        ]],
        colWidths=[4.5*cm, 4.5*cm, 4.5*cm, 4.5*cm],
        style=TableStyle([("ALIGN", (0,0), (-1,-1), "CENTER"), ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                          ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4)])
    )
    story.append(kpi_row)
    story.append(Spacer(1, 0.4*cm))

    # ── Sección 2: Parámetros del modelo ──────────────────────
    story.append(Paragraph(t["sec_model"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))
    param_data = [
        [t["param_cutoff"], cutoff_month],
        [t["param_actual"],  ", ".join(actual_months)],
        [t["param_forecast"], ", ".join(forecast_months)],
        [t["param_sens"],    f"{sensitivity_mult:.2f}x"],
        [t["param_mom"],     f"{momentum_mult:.2f}x"],
        [t["param_inf"],     f"{cagr_inf}%"],
        [t["param_ops"],     f"{cagr_ops}%"],
    ]
    param_rows = [[Paragraph(str(r[0]), s_body_bold), Paragraph(str(r[1]), s_body)] for r in param_data]
    tbl_params = make_table(
        [[Paragraph("Parámetro" if lang=="es" else "Parameter", s_body_bold),
          Paragraph("Valor" if lang=="es" else "Value", s_body_bold)]] + param_rows,
        [9*cm, 9*cm]
    )
    story.append(tbl_params)
    story.append(Spacer(1, 0.3*cm))

    # ── Sección 3: Resumen por categoría ──────────────────────
    story.append(Paragraph(t["sec_summary"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))

    sum_cols = [chart_dim, "Actual_YTD", "Budget_FY_Model", "Forecast_FY_Modelo", "Var_vs_Budget"]
    sum_show = summary[[c for c in sum_cols if c in summary.columns]].copy()

    header_row = [
        Paragraph(t["col_cat"], s_body_bold),
        Paragraph(t["col_actual"], s_body_bold),
        Paragraph(t["col_budget"], s_body_bold),
        Paragraph(t["col_forecast"], s_body_bold),
        Paragraph(t["col_var"], s_body_bold),
    ]
    data_rows = [header_row]
    for _, row in sum_show.iterrows():
        var_v = float(row.get("Var_vs_Budget", 0))
        var_color = colors.HexColor("#C00000") if var_v > 0 else (colors.HexColor("#375623") if var_v < 0 else colors.black)
        data_rows.append([
            Paragraph(str(row.get(chart_dim, "")), s_body),
            Paragraph(money(row.get("Actual_YTD", 0)), s_body),
            Paragraph(money(row.get("Budget_FY_Model", 0)), s_body),
            Paragraph(money(row.get("Forecast_FY_Modelo", 0)), s_body),
            Paragraph(f'<font color="#{("C00000" if var_v>0 else "375623" if var_v<0 else "000000")}">{money(var_v)}</font>', s_body),
        ])

    tbl_sum = make_table(data_rows, [4.5*cm, 3.5*cm, 3.5*cm, 3.5*cm, 3.5*cm])
    story.append(tbl_sum)
    story.append(Spacer(1, 0.3*cm))

    # ── Sección 4: LRP 5 años ─────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph(t["sec_lrp"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))

    lrp_num_cols = [c for c in resumen_5y.columns if c not in ["Contexto_Mina", "Naturaleza"]]
    lrp_dim_cols = [c for c in ["Contexto_Mina", "Naturaleza"] if c in resumen_5y.columns]
    all_lrp_cols = lrp_dim_cols + lrp_num_cols

    if len(all_lrp_cols) > 0 and len(resumen_5y) > 0:
        lrp_show = resumen_5y[all_lrp_cols].copy()
        # Mostrar solo FY cols (no meses individuales para que quepa en página)
        fy_cols_only = [c for c in lrp_num_cols if c.startswith("FY")]
        cols_to_show_lrp = lrp_dim_cols + fy_cols_only
        if cols_to_show_lrp:
            lrp_show = lrp_show[cols_to_show_lrp]
            n_cols = len(lrp_show.columns)
            dim_w  = 4.5*cm
            num_w  = (18*cm - dim_w * len(lrp_dim_cols)) / max(len(fy_cols_only), 1)
            lrp_header = [Paragraph(str(c), s_body_bold) for c in lrp_show.columns]
            lrp_rows = [lrp_header]
            for _, row in lrp_show.iterrows():
                lrp_row = []
                for i, (col, val) in enumerate(row.items()):
                    if col in lrp_dim_cols:
                        lrp_row.append(Paragraph(str(val), s_body_bold))
                    else:
                        try:
                            lrp_row.append(Paragraph(f"{float(val)/1_000_000:,.2f}", s_body))
                        except Exception:
                            lrp_row.append(Paragraph(str(val), s_body))
                lrp_rows.append(lrp_row)
            col_widths_lrp = [dim_w]*len(lrp_dim_cols) + [num_w]*len(fy_cols_only)
            story.append(make_table(lrp_rows, col_widths_lrp))
    story.append(Spacer(1, 0.3*cm))

    # ── Sección 5: Hallazgos y recomendaciones ─────────────────
    story.append(Paragraph(t["sec_hallazgos"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))
    story.append(Paragraph(t["hallazgo_1"], s_body))
    story.append(Paragraph(t["hallazgo_2"], s_body))
    story.append(Paragraph(t["hallazgo_3"], s_body))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(t["methodology"], s_h2))
    story.append(Paragraph(t["method_body"], s_body))
    story.append(Spacer(1, 0.2*cm))
    for rec in [t["rec_1"], t["rec_2"], t["rec_3"]]:
        story.append(Paragraph(f"• {rec}", s_body))
    story.append(Spacer(1, 0.3*cm))

    # ── Sección 6: Registro de sensibilidades ─────────────────
    story.append(Paragraph(t["sec_sens"], s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_LIGHT, spaceAfter=6))
    if not registro:
        story.append(Paragraph(t["no_sens"], s_caption))
    else:
        reg_cols_show = ["Escenario", "Timestamp", "Var Diesel", "Var FX", "Var MO",
                         "Budget Base LRP (US$ MM)", "Impacto Sens. (US$ MM)", "Budget Ajustado LRP (US$ MM)", "Var % vs Base"]
        df_reg = pd.DataFrame(registro)
        cols_avail = [c for c in reg_cols_show if c in df_reg.columns]
        df_reg_show = df_reg[cols_avail]
        reg_header = [Paragraph(str(c), s_body_bold) for c in df_reg_show.columns]
        reg_rows = [reg_header]
        col_w_reg = [18*cm / len(cols_avail)] * len(cols_avail)
        for _, row in df_reg_show.iterrows():
            reg_rows.append([Paragraph(str(v), s_body) for v in row.values])
        story.append(make_table(reg_rows, col_w_reg))

    # Footer en última página
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_ACCENT))
    story.append(Paragraph(t["footer"], s_caption))

    doc.build(story)
    return buf.getvalue()


def generar_docx_informe(
    lang: str,
    actual_total: float,
    budget_total: float,
    forecast_total: float,
    var_total: float,
    var_pct: float,
    summary: pd.DataFrame,
    chart_dim: str,
    actual_months: List[str],
    forecast_months: List[str],
    cutoff_month: str,
    sensitivity_mult: float,
    momentum_mult: float,
    cagr_inf: float,
    cagr_ops: float,
    resumen_5y: pd.DataFrame,
    registro: List[dict],
) -> bytes:
    """Genera informe ejecutivo en DOCX usando python-docx (sin Node.js)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE_DARK  = RGBColor(0x1F, 0x4E, 0x79)
    BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)
    BLUE_LIGHT = "D9EAF7"
    GRAY_BG    = "F2F2F2"
    RED_COLOR  = RGBColor(0xC0, 0x00, 0x00)
    GREEN_COLOR = RGBColor(0x37, 0x56, 0x23)
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

    T = {
        "es": {
            "title": "Informe Ejecutivo — Forecast 5+7 No Lineal Dinámico",
            "subtitle": "Análisis de Presupuesto y Proyección Minera",
            "date_label": "Fecha de generación",
            "period_label": "Período analizado",
            "period": f"Real: Jan — {cutoff_month} | Forecast: {forecast_months[0] if forecast_months else '—'} — {forecast_months[-1] if forecast_months else '—'}",
            "sec_kpi": "1. KPIs Ejecutivos",
            "sec_model": "2. Parámetros del Modelo",
            "sec_summary": "3. Resumen por Categoría",
            "sec_lrp": "4. Simulación LRP 5 Años (US$ MM)",
            "sec_hallazgos": "5. Hallazgos y Recomendaciones",
            "sec_sens": "6. Registro de Sensibilidades",
            "col_cat": "Categoría", "col_actual": "Actual YTD", "col_budget": "Budget FY",
            "col_forecast": "Forecast FY", "col_var": "Var vs Budget",
            "param_label": "Parámetro", "value_label": "Valor",
            "param_cutoff": "Último mes real", "param_actual": "Meses reales",
            "param_forecast": "Meses forecast", "param_sens": "Sensibilidad ejecución",
            "param_mom": "Peso tendencia", "param_inf": "Inflación anual",
            "param_ops": "Crec. operacional",
            "kpi_actual": "Actual YTD", "kpi_budget": "Budget FY",
            "kpi_forecast": "Forecast FY Modelo", "kpi_var": "Variación vs Budget",
            "hallazgo_1": f"El Forecast FY Modelo proyecta {money(forecast_total)}, con desviación {var_pct:+.1%} vs Budget FY de {money(budget_total)}.",
            "hallazgo_2": "Las categorías con mayor peso presupuestario concentran las principales desviaciones y deben priorizarse en el control operacional.",
            "hallazgo_3": f"Simulación LRP 5 años con inflación {cagr_inf}% y crecimiento operacional {cagr_ops}%.",
            "rec_1": "Activar plan de control en categorías con desviación positiva superior al 10%.",
            "rec_2": "Revisar reprogramación en partidas con subejecución relevante.",
            "rec_3": "Implementar alertas tempranas mensuales por naturaleza de gasto.",
            "footer": "Modelo Forecast 5+7 No Lineal Dinámico — Uso interno de gestión presupuestaria",
            "no_sens": "No se registraron escenarios de sensibilidad en esta sesión.",
            "findings_intro": "El modelo Forecast 5+7 proyecta el cierre anual combinando ejecución real acumulada, presupuesto restante y una curva no lineal ajustada al comportamiento reciente del gasto.",
        },
        "en": {
            "title": "Executive Report — Dynamic Non-Linear 5+7 Forecast",
            "subtitle": "Mining Budget Analysis & Projection",
            "date_label": "Generated on",
            "period_label": "Analysis period",
            "period": f"Actual: Jan — {cutoff_month} | Forecast: {forecast_months[0] if forecast_months else '—'} — {forecast_months[-1] if forecast_months else '—'}",
            "sec_kpi": "1. Executive KPIs",
            "sec_model": "2. Model Parameters",
            "sec_summary": "3. Summary by Category",
            "sec_lrp": "4. 5-Year LRP Simulation (US$ MM)",
            "sec_hallazgos": "5. Findings & Recommendations",
            "sec_sens": "6. Sensitivity Scenarios Log",
            "col_cat": "Category", "col_actual": "Actual YTD", "col_budget": "Budget FY",
            "col_forecast": "Forecast FY", "col_var": "Var vs Budget",
            "param_label": "Parameter", "value_label": "Value",
            "param_cutoff": "Last actual month", "param_actual": "Actual months",
            "param_forecast": "Forecast months", "param_sens": "Execution sensitivity",
            "param_mom": "Trend weight", "param_inf": "Annual inflation",
            "param_ops": "Ops growth",
            "kpi_actual": "Actual YTD", "kpi_budget": "Budget FY",
            "kpi_forecast": "Forecast FY Model", "kpi_var": "Variance vs Budget",
            "hallazgo_1": f"FY Model Forecast projects {money(forecast_total)}, {var_pct:+.1%} deviation vs FY Budget of {money(budget_total)}.",
            "hallazgo_2": "Categories with the highest budget weight concentrate the main deviations and should be prioritized in operational control.",
            "hallazgo_3": f"5-year LRP with {cagr_inf}% inflation and {cagr_ops}% operational growth.",
            "rec_1": "Activate control plan for categories with positive deviation above 10%.",
            "rec_2": "Review rescheduling for categories with significant under-execution.",
            "rec_3": "Implement monthly early-warning alerts by cost nature.",
            "footer": "Dynamic Non-Linear 5+7 Forecast Model — Internal budget management use",
            "no_sens": "No sensitivity scenarios were recorded in this session.",
            "findings_intro": "The 5+7 Forecast model projects the annual close combining actual YTD spend, remaining budget, and a non-linear curve adjusted to recent spending behavior.",
        },
    }
    t = T[lang]

    doc = Document()

    # ── Estilos globales ────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_heading(text, level=1, color=None):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.color.rgb = color or (BLUE_DARK if level == 1 else BLUE_MID)
        run.font.size = Pt(14 if level == 1 else 11)
        p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
        p.paragraph_format.space_after = Pt(4)
        if level == 1:
            p.paragraph_format.border_bottom = None
        return p

    def add_body(text, bold=False, color=None, indent=False):
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run(str(text))
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.8)
        return p

    def add_hr():
        p = doc.add_paragraph()
        p.clear()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)

    def add_table(headers, rows, col_widths_cm, alternate=True):
        n_cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = BLUE_DARK
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(hdr_cells[i], BLUE_LIGHT)
            tbl.columns[i].width = Cm(col_widths_cm[i])

        # Data rows
        for r_idx, row in enumerate(rows):
            cells = tbl.rows[r_idx + 1].cells
            bg = "FFFFFF" if (r_idx % 2 == 0 or not alternate) else GRAY_BG
            for c_idx, val in enumerate(row):
                cells[c_idx].text = str(val)
                run = cells[c_idx].paragraphs[0].runs[0] if cells[c_idx].paragraphs[0].runs else cells[c_idx].paragraphs[0].add_run(str(val))
                run.font.name = "Arial"
                run.font.size = Pt(8)
                cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                set_cell_bg(cells[c_idx], bg)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        return tbl

    # ── PORTADA ─────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(t["title"])
    run_t.font.name = "Arial"
    run_t.font.bold = True
    run_t.font.size = Pt(18)
    run_t.font.color.rgb = BLUE_DARK

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(t["subtitle"])
    run_s.font.name = "Arial"
    run_s.font.size = Pt(12)
    run_s.font.color.rgb = BLUE_MID

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_date.add_run(f"{t['date_label']}: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run_d.font.name = "Arial"
    run_d.font.size = Pt(9)
    run_d.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p_per = doc.add_paragraph()
    p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_p = p_per.add_run(f"{t['period_label']}: {t['period']}")
    run_p.font.name = "Arial"
    run_p.font.size = Pt(9)
    run_p.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_hr()

    # ── SECCIÓN 1: KPIs ─────────────────────────────────────────
    add_heading(t["sec_kpi"], level=1)
    add_hr()

    kpi_tbl = doc.add_table(rows=2, cols=4)
    kpi_tbl.style = "Table Grid"
    kpi_headers = [t["kpi_actual"], t["kpi_budget"], t["kpi_forecast"], t["kpi_var"]]
    kpi_values  = [money(actual_total), money(budget_total), money(forecast_total),
                   f"{var_pct:+.1%} ({money(var_total)})"]
    kpi_bg = ["2E75B6", "2E75B6", "2E75B6", "C00000" if var_total > 0 else "375623"]

    for i in range(4):
        lbl_cell = kpi_tbl.rows[0].cells[i]
        val_cell = kpi_tbl.rows[1].cells[i]
        lbl_cell.text = kpi_headers[i]
        val_cell.text = kpi_values[i]
        for cell in [lbl_cell, val_cell]:
            set_cell_bg(cell, kpi_bg[i])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.color.rgb = WHITE
        lbl_cell.paragraphs[0].runs[0].font.size = Pt(9)
        val_cell.paragraphs[0].runs[0].font.size = Pt(12)
        val_cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── SECCIÓN 2: Parámetros ───────────────────────────────────
    add_heading(t["sec_model"], level=1)
    add_hr()
    params_data = [
        [t["param_cutoff"],   cutoff_month],
        [t["param_actual"],   ", ".join(actual_months)],
        [t["param_forecast"], ", ".join(forecast_months)],
        [t["param_sens"],     f"{sensitivity_mult:.2f}x"],
        [t["param_mom"],      f"{momentum_mult:.2f}x"],
        [t["param_inf"],      f"{cagr_inf}%"],
        [t["param_ops"],      f"{cagr_ops}%"],
    ]
    add_table([t["param_label"], t["value_label"]], params_data, [8.5, 8.5])

    # ── SECCIÓN 3: Resumen por categoría ────────────────────────
    add_heading(t["sec_summary"], level=1)
    add_hr()
    sum_cols = [chart_dim, "Actual_YTD", "Budget_FY_Model", "Forecast_FY_Modelo", "Var_vs_Budget"]
    sum_show = summary[[c for c in sum_cols if c in summary.columns]].copy()
    sum_headers = [t["col_cat"], t["col_actual"], t["col_budget"], t["col_forecast"], t["col_var"]]
    sum_rows = [
        [str(row.get(chart_dim, "")), money(row.get("Actual_YTD", 0)),
         money(row.get("Budget_FY_Model", 0)), money(row.get("Forecast_FY_Modelo", 0)),
         money(row.get("Var_vs_Budget", 0))]
        for _, row in sum_show.iterrows()
    ]
    add_table(sum_headers, sum_rows, [4.0, 3.1, 3.1, 3.1, 3.1])

    # ── SECCIÓN 4: LRP 5 años ───────────────────────────────────
    doc.add_page_break()
    add_heading(t["sec_lrp"], level=1)
    add_hr()
    fy_cols_lrp  = [c for c in resumen_5y.columns if c.startswith("FY")]
    lrp_dim_cols = [c for c in ["Contexto_Mina", "Naturaleza"] if c in resumen_5y.columns]
    lrp_show_cols = lrp_dim_cols + fy_cols_lrp
    lrp_show = resumen_5y[[c for c in lrp_show_cols if c in resumen_5y.columns]].copy()

    if len(lrp_show) > 0 and len(fy_cols_lrp) > 0:
        n_dim = len(lrp_dim_cols)
        n_fy  = len(fy_cols_lrp)
        dim_w = 4.0
        fy_w  = max(1.5, (17.0 - dim_w * n_dim) / max(n_fy, 1))
        col_w = [dim_w] * n_dim + [fy_w] * n_fy
        lrp_rows = []
        for _, row in lrp_show.iterrows():
            r = []
            for col, val in row.items():
                if col in lrp_dim_cols:
                    r.append(str(val))
                else:
                    try:
                        r.append(f"{float(val)/1_000_000:,.2f}")
                    except Exception:
                        r.append(str(val))
            lrp_rows.append(r)
        add_table(list(lrp_show.columns), lrp_rows, col_w)
    else:
        add_body("Sin datos LRP disponibles." if lang == "es" else "No LRP data available.")

    # ── SECCIÓN 5: Hallazgos ────────────────────────────────────
    add_heading(t["sec_hallazgos"], level=1)
    add_hr()
    add_body(t["findings_intro"])
    doc.add_paragraph()
    for txt in [t["hallazgo_1"], t["hallazgo_2"], t["hallazgo_3"]]:
        add_body(f"• {txt}", indent=True)
    doc.add_paragraph()
    add_heading("Recomendaciones" if lang == "es" else "Recommendations", level=2)
    for rec in [t["rec_1"], t["rec_2"], t["rec_3"]]:
        add_body(f"• {rec}", indent=True)

    # ── SECCIÓN 6: Registro sensibilidades ─────────────────────
    add_heading(t["sec_sens"], level=1)
    add_hr()
    if not registro:
        add_body(t["no_sens"], color=RGBColor(0x66, 0x66, 0x66))
    else:
        reg_cols_show = ["Escenario", "Timestamp", "Var Diesel", "Var FX", "Var MO",
                         "Budget Base LRP (US$ MM)", "Impacto Sens. (US$ MM)",
                         "Budget Ajustado LRP (US$ MM)", "Var % vs Base"]
        df_reg = pd.DataFrame(registro)
        cols_avail = [c for c in reg_cols_show if c in df_reg.columns]
        if cols_avail:
            reg_rows = [[str(df_reg.loc[i, c]) for c in cols_avail] for i in df_reg.index]
            n = len(cols_avail)
            reg_col_w = [17.0 / n] * n
            add_table(cols_avail, reg_rows, reg_col_w)

    # ── Footer ──────────────────────────────────────────────────
    doc.add_paragraph()
    add_hr()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run(t["footer"])
    run_f.font.name = "Arial"
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ═══════════════════════════════════════════════════════════════
# ESTADO DE SESIÓN PARA REGISTRO DE SENSIBILIDADES
# ═══════════════════════════════════════════════════════════════
if "registro_sensibilidades" not in st.session_state:
    st.session_state["registro_sensibilidades"] = []

# ═══════════════════════════════════════════════════════════════
# PESTAÑAS PRINCIPALES
# ═══════════════════════════════════════════════════════════════
tab_actual, tab_5y, tab_sens, tab_registro, tab_informe = st.tabs([
    "📊 Forecast Año Actual (5+7)",
    "🚀 Simulación LRP a 5 Años",
    "⚙️ Módulo de Sensibilidades",
    "📋 Registro de Versiones",
    "📄 Informe Ejecutivo",
])

# ─────────────────────────────────────────────
# PESTAÑA 1: FORECAST AÑO ACTUAL
# ─────────────────────────────────────────────
with tab_actual:

    c1, c2 = st.columns(2)
    with c1:
        fig = px.bar(summary.head(15), x=chart_dim, y=["Budget_FY_Model", "Forecast_FY_Modelo"], barmode="group", title="Forecast FY vs Budget FY")
        fig.update_layout(xaxis_title="Categoría", yaxis_title="Monto")
        st.plotly_chart(fig, use_container_width=True)
    with c2:
        fig2 = px.bar(summary.sort_values("Var_vs_Budget", ascending=False).head(15), x=chart_dim, y="Var_vs_Budget", title="Principales desviaciones vs Budget")
        fig2.update_layout(xaxis_title="Categoría", yaxis_title="Varianza")
        st.plotly_chart(fig2, use_container_width=True)

    # Curva mensual total
    monthly_rows = []
    for m in MONTH_ORDER:
        if m in actual_months:
            monthly_rows.append({"Mes": m, "Tipo": "Actual", "Monto": filtered.get(f"Actual_{m}", pd.Series(0, index=filtered.index)).sum()})
        elif m in forecast_months:
            monthly_rows.append({"Mes": m, "Tipo": "Forecast Modelo", "Monto": filtered.get(f"Forecast_{m}", pd.Series(0, index=filtered.index)).sum()})
        monthly_rows.append({"Mes": m, "Tipo": "Budget", "Monto": filtered.get(f"Budget_{m}", pd.Series(0, index=filtered.index)).sum()})
    monthly_df = pd.DataFrame(monthly_rows)
    fig3 = px.line(monthly_df, x="Mes", y="Monto", color="Tipo", markers=True, title="Serie mensual: Actual + Forecast no lineal vs Budget")
    st.plotly_chart(fig3, use_container_width=True)

    st.subheader("📌 Hallazgos automáticos")
    if len(summary) > 0:
        top_spend = summary.iloc[0]
        top_over = summary.sort_values("Var_vs_Budget", ascending=False).iloc[0]
        top_under = summary.sort_values("Var_vs_Budget", ascending=True).iloc[0]
        st.markdown(f"""
        - La mayor concentración de gasto proyectado está en **{top_spend[chart_dim]}**, con un forecast de **{money(top_spend['Forecast_FY_Modelo'])}**.
        - La mayor desviación positiva contra presupuesto está en **{top_over[chart_dim]}**, con **{money(top_over['Var_vs_Budget'])}** sobre Budget.
        - La mayor subejecución proyectada está en **{top_under[chart_dim]}**, con **{money(top_under['Var_vs_Budget'])}** respecto al Budget.
        - El forecast total del filtro seleccionado queda en **{money(forecast_total)}**, equivalente a una desviación de **{var_pct:.1%}** contra el presupuesto anual.
        """)

    # Gráfico Waterfall
    st.subheader("🌊 Gráfico cascada: explicación de la variación presupuestaria")
    waterfall_df = summary.copy()
    top_over_wf = waterfall_df[waterfall_df["Var_vs_Budget"] > 0].sort_values("Var_vs_Budget", ascending=False).head(3)
    top_under_wf = waterfall_df[waterfall_df["Var_vs_Budget"] < 0].sort_values("Var_vs_Budget", ascending=True).head(3)
    waterfall_items = pd.concat([top_over_wf, top_under_wf])

    inicio = budget_total
    final = forecast_total
    labels = ["Budget FY"] + waterfall_items[chart_dim].astype(str).tolist() + ["Forecast FY"]
    values = [inicio] + waterfall_items["Var_vs_Budget"].tolist() + [final]
    measure = ["absolute"] + ["relative"] * len(waterfall_items) + ["total"]

    fig_waterfall = go.Figure(go.Waterfall(
        name="Variación", orientation="v", measure=measure, x=labels, y=values,
        text=[money(inicio)] + [money(v) for v in waterfall_items["Var_vs_Budget"]] + [money(final)],
        textposition="outside",
        connector={"line": {"color": "gray"}},
        increasing={"marker": {"color": "#4C78A8"}},
        decreasing={"marker": {"color": "#E45756"}},
        totals={"marker": {"color": "#72B7B2"}},
    ))
    fig_waterfall.update_layout(
        title="Explicación de la variación entre Budget FY y Forecast FY",
        yaxis_title="Monto", xaxis_title="Categoría", showlegend=False, height=520,
    )
    st.plotly_chart(fig_waterfall, use_container_width=True)

    if len(waterfall_items) > 0:
        top_pos = waterfall_items[waterfall_items["Var_vs_Budget"] > 0].sort_values("Var_vs_Budget", ascending=False)
        top_neg = waterfall_items[waterfall_items["Var_vs_Budget"] < 0].sort_values("Var_vs_Budget", ascending=True)
        texto_pos = f"- La mayor desviación positiva corresponde a **{top_pos.iloc[0][chart_dim]}**, con un aumento de **{money(top_pos.iloc[0]['Var_vs_Budget'])}** respecto al presupuesto.\n" if len(top_pos) > 0 else "- No se observan categorías con desviación positiva relevante.\n"
        texto_neg = f"- La mayor desviación negativa corresponde a **{top_neg.iloc[0][chart_dim]}**, con una reducción de **{money(top_neg.iloc[0]['Var_vs_Budget'])}** respecto al presupuesto.\n" if len(top_neg) > 0 else "- No se observan categorías con desviación negativa relevante.\n"
        conclusion_var = "El modelo proyecta una **subejecución presupuestaria**." if var_total < 0 else ("El modelo proyecta una **sobreejecución presupuestaria**." if var_total > 0 else "El modelo proyecta un cierre alineado con el presupuesto.")
        st.markdown(f"""
### Interpretación automática
- Budget FY: **{money(budget_total)}** | Forecast FY Modelo: **{money(forecast_total)}** | Desviación: **{money(var_total)} ({var_pct:.1%})**
{texto_pos}{texto_neg}{conclusion_var}
        """)

    st.subheader("✅ Propuesta formal de mejora")
    st.markdown("Se recomienda implementar un sistema de control presupuestario con alertas tempranas por naturaleza de gasto y centro de costo. Las partidas con desviaciones positivas deben revisarse mediante acciones de control operacional, renegociación contractual y revisión de consumos críticos.")

    st.subheader("📄 Resultado detallado")
    cols_to_show = [c for c in extra_dims if c in filtered.columns] + ["Naturaleza", "Contexto_Mina", "Actual_YTD", "Budget_YTD", "Budget_Remaining", "Forecast_Remaining", "Budget_FY_Model", "Forecast_FY_Modelo", "Var_vs_Budget", "Var_vs_Budget_%", "Recomendacion", "Justificacion_Mina"]
    cols_to_show = list(dict.fromkeys([c for c in cols_to_show if c in filtered.columns]))
    st.dataframe(filtered[cols_to_show].sort_values("Var_vs_Budget", ascending=False), use_container_width=True, height=420)

    excel_bytes = to_excel_bytes(filtered[cols_to_show], summary)
    st.download_button(
        "⬇️ Descargar Forecast en Excel",
        data=excel_bytes,
        file_name="forecast_5mas7_no_lineal_resultados.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# ─────────────────────────────────────────────
# PESTAÑA 2: SIMULACIÓN LRP 5 AÑOS
# ─────────────────────────────────────────────
with tab_5y:
    año_base = datetime.datetime.now().year

    st.subheader("Simulación a 5 Años basada en Sensibilidad Operativa")
    st.caption(f"Inflación: {cagr_inf}% | Crecimiento operacional: {cagr_ops}% | Base: Forecast FY Modelo")

    if len(filtered) > 0:
        df_5y = simular_5_anos(filtered, cagr_inf, cagr_ops)

        cols_agrup = ["Contexto_Mina"]
        if "Naturaleza" in df_5y.columns:
            cols_agrup = ["Contexto_Mina", "Naturaleza"]

        año_cols = ["Año_0_Base", "Año_1", "Año_2", "Año_3", "Año_4", "Año_5"]
        resumen_5y_raw = df_5y.groupby(cols_agrup)[año_cols].sum().reset_index()

        # Estacionalidad para desglosar Año_1 en 12 meses
        estacionalidad = calcular_estacionalidad_mensual(filtered, actual_months, forecast_months)
        año1_label = año_base + 1
        meses_año1 = [f"{m}-{str(año1_label)[2:]}" for m in MONTH_ORDER]

        for m, mes_col in zip(MONTH_ORDER, meses_año1):
            resumen_5y_raw[mes_col] = resumen_5y_raw["Año_1"] * estacionalidad[m]

        # Renombrar años acumulados (Año_0_Base = FY26, Año_2..5 = FY28..31)
        rename_map = {
            "Año_0_Base": f"FY{str(año_base)[2:]}",
            "Año_2": f"FY{str(año_base + 2)[2:]}",
            "Año_3": f"FY{str(año_base + 3)[2:]}",
            "Año_4": f"FY{str(año_base + 4)[2:]}",
            "Año_5": f"FY{str(año_base + 5)[2:]}",
        }
        resumen_5y = resumen_5y_raw.rename(columns=rename_map).drop(columns=["Año_1"])

        fy_acumulados = [f"FY{str(año_base + i)[2:]}" for i in [0, 2, 3, 4, 5]]
        cols_orden = cols_agrup + [fy_acumulados[0]] + meses_año1 + fy_acumulados[1:]
        resumen_5y = resumen_5y[[c for c in cols_orden if c in resumen_5y.columns]]

        # Fila TOTAL
        cols_money = [c for c in resumen_5y.columns if c not in cols_agrup]
        total_row = {}
        for i, col in enumerate(resumen_5y.columns):
            if col in cols_money:
                total_row[col] = resumen_5y[col].sum()
            elif i == 0:
                total_row[col] = "TOTAL"
            else:
                total_row[col] = ""
        resumen_5y_con_total = pd.concat([resumen_5y, pd.DataFrame([total_row])], ignore_index=True)

        def fmt_mm(val):
            try:
                return f"US$ {float(val)/1_000_000:,.2f} MM"
            except Exception:
                return val

        resumen_display = resumen_5y_con_total.copy()
        for col in cols_money:
            resumen_display[col] = resumen_display[col].apply(fmt_mm)

        st.markdown(f"**Proyección LRP — {año1_label} mensual + acumulados anuales (millones USD)**")
        st.dataframe(resumen_display, use_container_width=True, hide_index=True)

        # Gráfico de áreas (sin FY26 base, arranca en FY27)
        año_cols_sin_base = ["Año_1", "Año_2", "Año_3", "Año_4", "Año_5"]
        año_cols_grafico  = [f"FY{str(año_base + i)[2:]}" for i in range(1, 6)]
        resumen_grafico = df_5y.groupby(cols_agrup)[año_cols_sin_base].sum().reset_index()
        resumen_grafico = resumen_grafico.rename(columns=dict(zip(año_cols_sin_base, año_cols_grafico)))
        resumen_largo = resumen_grafico.melt(
            id_vars=cols_agrup, var_name="Año", value_name="Presupuesto Proyectado"
        )

        fig_5y = px.area(
            resumen_largo,
            x="Año",
            y="Presupuesto Proyectado",
            color="Contexto_Mina",
            title=f"Evolución Estructural del Presupuesto FY{str(año_base+1)[2:]}–FY{str(año_base+5)[2:]}",
            markers=True,
        )
        st.plotly_chart(fig_5y, use_container_width=True)

        # Botón descarga reporte completo
        excel_completo = to_excel_completo(
            filtered, summary, resumen_5y, resumen_largo, actual_months, forecast_months
        )
        st.download_button(
            label="⬇️ Descargar Reporte Completo: Forecast 5+7 + Simulación LRP (Excel)",
            data=excel_completo,
            file_name="reporte_forecast_lrp.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    else:
        st.warning("⚠️ No hay datos para simular. Ajusta los filtros en el menú lateral.")

# ─────────────────────────────────────────────
# PESTAÑA 3: MÓDULO DE SENSIBILIDADES
# ─────────────────────────────────────────────
with tab_sens:
    st.subheader("⚙️ Módulo de Sensibilidades Macroeconómicas")
    st.caption("Analiza el impacto económico de variaciones en precios de combustible, divisa y mano de obra sobre el Budget 2027–2031.")

    if len(filtered) == 0:
        st.warning("⚠️ No hay datos cargados. Ajusta los filtros en el menú lateral.")
    else:
        año_base_s = datetime.datetime.now().year

        # ── Calcular base LRP sin sensibilidades ──────────────────────────
        df_base_s = simular_5_anos(filtered, cagr_inf, cagr_ops)
        cols_agrup_s = ["Contexto_Mina"]
        año_cols_s = ["Año_1", "Año_2", "Año_3", "Año_4", "Año_5"]
        base_lrp = df_base_s.groupby(cols_agrup_s)[año_cols_s].sum()
        total_base_por_año = base_lrp.sum()  # Serie: Año_1..5

        # Pesos de cada categoría sobre el total base (para calcular impacto parcial)
        peso_fuel    = base_lrp.loc["Fuel"].sum()    / total_base_por_año.sum() if "Fuel"  in base_lrp.index else 0.0
        peso_power   = base_lrp.loc["Power"].sum()   / total_base_por_año.sum() if "Power" in base_lrp.index else 0.0
        peso_labor   = base_lrp.loc["Labor"].sum()   / total_base_por_año.sum() if "Labor" in base_lrp.index else 0.0

        st.markdown("---")
        st.markdown("### 🔧 Configurar escenario de sensibilidad")

        nombre_escenario = st.text_input("Nombre del escenario", value=f"Escenario {len(st.session_state['registro_sensibilidades'])+1}")

        col_s1, col_s2, col_s3 = st.columns(3)

        with col_s1:
            st.markdown("#### ⛽ Combustible / Diesel")
            precio_diesel_base = st.number_input("Precio base diesel (US$/lt)", value=0.85, step=0.01, format="%.3f")
            precio_diesel_nuevo = st.number_input("Precio estimado diesel (US$/lt)", value=0.85, step=0.01, format="%.3f")
            var_diesel_pct = safe_div(precio_diesel_nuevo - precio_diesel_base, precio_diesel_base, 0.0)
            st.metric("Variación Diesel", f"{var_diesel_pct:+.1%}", delta_color="inverse")

        with col_s2:
            st.markdown("#### 💱 Divisa (CLP/USD)")
            tipo_cambio_base = st.number_input("Tipo de cambio base (CLP/USD)", value=900.0, step=1.0, format="%.0f")
            tipo_cambio_nuevo = st.number_input("Tipo de cambio estimado (CLP/USD)", value=900.0, step=1.0, format="%.0f")
            var_fx_pct = safe_div(tipo_cambio_nuevo - tipo_cambio_base, tipo_cambio_base, 0.0)
            st.metric("Variación Divisa", f"{var_fx_pct:+.1%}", delta_color="normal")

        with col_s3:
            st.markdown("#### 👷 Mano de Obra (Labor)")
            costo_mo_base = st.number_input("Costo MO base (US$/hora)", value=25.0, step=0.5, format="%.2f")
            costo_mo_nuevo = st.number_input("Costo MO estimado (US$/hora)", value=25.0, step=0.5, format="%.2f")
            var_mo_pct = safe_div(costo_mo_nuevo - costo_mo_base, costo_mo_base, 0.0)
            st.metric("Variación Mano de Obra", f"{var_mo_pct:+.1%}", delta_color="inverse")

        st.markdown("---")

        # ── Parámetros de sensibilidad por categoría ───────────────────────
        with st.expander("🎚️ Ajustar sensibilidad por categoría (avanzado)", expanded=False):
            st.caption("Define qué % del costo de cada categoría es sensible a cada variable macro.")
            cs1, cs2, cs3 = st.columns(3)
            with cs1:
                sens_diesel_fuel  = st.slider("Sensibilidad Diesel → Fuel (%)", 0, 100, 80) / 100
                sens_diesel_spare = st.slider("Sensibilidad Diesel → Spare Parts (%)", 0, 100, 20) / 100
                sens_diesel_maint = st.slider("Sensibilidad Diesel → Maintenance (%)", 0, 100, 10) / 100
            with cs2:
                sens_fx_fuel      = st.slider("Sensibilidad FX → Fuel (%)", 0, 100, 60) / 100
                sens_fx_spare     = st.slider("Sensibilidad FX → Spare Parts (%)", 0, 100, 50) / 100
                sens_fx_contract  = st.slider("Sensibilidad FX → Contractors (%)", 0, 100, 30) / 100
            with cs3:
                sens_mo_labor     = st.slider("Sensibilidad MO → Labor (%)", 0, 100, 90) / 100
                sens_mo_contract  = st.slider("Sensibilidad MO → Contractors (%)", 0, 100, 40) / 100
                sens_mo_maint     = st.slider("Sensibilidad MO → Maintenance (%)", 0, 100, 20) / 100

        # ── Calcular impacto por año y categoría ──────────────────────────
        def calcular_impacto_sens(base_lrp, var_diesel, var_fx, var_mo,
                                   sd_fuel, sd_spare, sd_maint,
                                   sfx_fuel, sfx_spare, sfx_contract,
                                   smo_labor, smo_contract, smo_maint):
            impacto = {}
            for ctx in base_lrp.index:
                row = base_lrp.loc[ctx].copy()
                delta = pd.Series(0.0, index=row.index)
                if ctx == "Fuel":
                    delta += row * (var_diesel * sd_fuel + var_fx * sfx_fuel)
                elif ctx == "Spare Parts":
                    delta += row * (var_diesel * sd_spare + var_fx * sfx_spare)
                elif ctx == "Maintenance":
                    delta += row * (var_diesel * sd_maint + var_mo * smo_maint)
                elif ctx == "Labor":
                    delta += row * (var_mo * smo_labor)
                elif ctx == "Contractors":
                    delta += row * (var_fx * sfx_contract + var_mo * smo_contract)
                impacto[ctx] = delta
            return pd.DataFrame(impacto).T

        df_impacto = calcular_impacto_sens(
            base_lrp, var_diesel_pct, var_fx_pct, var_mo_pct,
            sens_diesel_fuel, sens_diesel_spare, sens_diesel_maint,
            sens_fx_fuel, sens_fx_spare, sens_fx_contract,
            sens_mo_labor, sens_mo_contract, sens_mo_maint,
        )

        total_base   = total_base_por_año.sum()
        total_impacto_por_año = df_impacto.sum()
        total_impacto = total_impacto_por_año.sum()
        total_nuevo  = total_base + total_impacto
        var_pct_sens = safe_div(total_impacto, total_base, 0.0)

        # ── KPIs de impacto ────────────────────────────────────────────────
        st.markdown("### 📊 Impacto estimado sobre Budget LRP 2027–2031")
        ks1, ks2, ks3, ks4 = st.columns(4)
        ks1.metric("Budget Base LRP", money(total_base))
        ks2.metric("Impacto Total Sensibilidades", money(total_impacto), f"{var_pct_sens:+.1%}", delta_color="inverse")
        ks3.metric("Budget Ajustado LRP", money(total_nuevo))
        ks4.metric("Variación vs Base", f"{var_pct_sens:+.1%}", delta_color="inverse")

        # ── Tabla impacto por categoría y año ─────────────────────────────
        st.markdown("#### Impacto por categoría (US$ MM)")
        fy_labels = [f"FY{str(año_base_s + i)[2:]}" for i in range(1, 6)]
        df_impacto_display = df_impacto.copy()
        df_impacto_display.columns = fy_labels
        df_impacto_display["Total Impacto"] = df_impacto_display.sum(axis=1)
        df_impacto_display["% sobre Base"] = [
            safe_div(df_impacto_display.loc[ctx, "Total Impacto"], base_lrp.loc[ctx].sum(), 0.0)
            if ctx in base_lrp.index else 0.0
            for ctx in df_impacto_display.index
        ]
        df_impacto_display_fmt = df_impacto_display.copy()
        for col in fy_labels + ["Total Impacto"]:
            df_impacto_display_fmt[col] = df_impacto_display_fmt[col].apply(lambda v: f"US$ {v/1_000_000:,.2f} MM")
        df_impacto_display_fmt["% sobre Base"] = df_impacto_display_fmt["% sobre Base"].apply(lambda v: f"{v:+.1%}")
        st.dataframe(df_impacto_display_fmt, use_container_width=True)

        # ── Gráfico tornado ────────────────────────────────────────────────
        st.markdown("#### 🌪️ Gráfico Tornado — Impacto por categoría")
        tornado_df = pd.DataFrame({
            "Categoría": df_impacto_display.index,
            "Impacto (US$ MM)": df_impacto_display["Total Impacto"] / 1_000_000,
        }).sort_values("Impacto (US$ MM)")
        fig_tornado = px.bar(
            tornado_df, x="Impacto (US$ MM)", y="Categoría", orientation="h",
            color="Impacto (US$ MM)",
            color_continuous_scale=["#E45756", "#FFFFFF", "#4C78A8"],
            color_continuous_midpoint=0,
            title="Impacto neto por categoría ante variaciones macro",
        )
        fig_tornado.update_layout(showlegend=False, height=400)
        st.plotly_chart(fig_tornado, use_container_width=True)

        # ── Gráfico evolución anual base vs ajustado ───────────────────────
        st.markdown("#### 📈 Evolución anual: Base vs Ajustado por sensibilidades")
        evol_df = pd.DataFrame({
            "Año": fy_labels,
            "Base": total_base_por_año.values / 1_000_000,
            "Ajustado": (total_base_por_año + total_impacto_por_año).values / 1_000_000,
        })
        fig_evol = go.Figure()
        fig_evol.add_trace(go.Scatter(x=evol_df["Año"], y=evol_df["Base"], name="Budget Base", mode="lines+markers", line=dict(color="#4C78A8", width=2)))
        fig_evol.add_trace(go.Scatter(x=evol_df["Año"], y=evol_df["Ajustado"], name="Budget Ajustado", mode="lines+markers", line=dict(color="#E45756", width=2, dash="dash")))
        fig_evol.update_layout(title="Budget Base vs Ajustado por sensibilidades (US$ MM)", yaxis_title="US$ MM", height=380)
        st.plotly_chart(fig_evol, use_container_width=True)

        # ── Botón guardar escenario en registro ───────────────────────────
        st.markdown("---")
        if st.button(f"💾 Guardar escenario '{nombre_escenario}' en el Registro", use_container_width=True):
            registro_entry = {
                "Escenario": nombre_escenario,
                "Timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                "Diesel base (US$/lt)": precio_diesel_base,
                "Diesel estimado (US$/lt)": precio_diesel_nuevo,
                "Var Diesel": f"{var_diesel_pct:+.1%}",
                "FX base (CLP/USD)": tipo_cambio_base,
                "FX estimado (CLP/USD)": tipo_cambio_nuevo,
                "Var FX": f"{var_fx_pct:+.1%}",
                "MO base (US$/hr)": costo_mo_base,
                "MO estimado (US$/hr)": costo_mo_nuevo,
                "Var MO": f"{var_mo_pct:+.1%}",
                "Inflación (%)": cagr_inf,
                "Crec. Ops (%)": cagr_ops,
                "Budget Base LRP (US$ MM)": round(total_base / 1_000_000, 2),
                "Impacto Sens. (US$ MM)": round(total_impacto / 1_000_000, 2),
                "Budget Ajustado LRP (US$ MM)": round(total_nuevo / 1_000_000, 2),
                "Var % vs Base": f"{var_pct_sens:+.1%}",
            }
            # Agregar impacto por año
            for fy, imp in zip(fy_labels, total_impacto_por_año.values):
                registro_entry[f"Impacto {fy} (US$ MM)"] = round(imp / 1_000_000, 2)
            # Agregar budget ajustado por año
            for fy, base, imp in zip(fy_labels, total_base_por_año.values, total_impacto_por_año.values):
                registro_entry[f"Ajustado {fy} (US$ MM)"] = round((base + imp) / 1_000_000, 2)

            st.session_state["registro_sensibilidades"].append(registro_entry)
            st.success(f"✅ Escenario '{nombre_escenario}' guardado. Ve a la pestaña 📋 Registro de Versiones para comparar.")


# ─────────────────────────────────────────────
# PESTAÑA 4: REGISTRO DE VERSIONES
# ─────────────────────────────────────────────
with tab_registro:
    st.subheader("📋 Registro de Versiones de Budget 2027–2031")
    st.caption("Historial de escenarios generados en esta sesión. Cada iteración del módulo de sensibilidades genera una nueva versión del Budget.")

    if len(st.session_state["registro_sensibilidades"]) == 0:
        st.info("Aún no hay versiones registradas. Ve al módulo ⚙️ Sensibilidades, configura un escenario y presiona 'Guardar escenario'.")
    else:
        df_registro = pd.DataFrame(st.session_state["registro_sensibilidades"])

        st.markdown(f"**{len(df_registro)} versión(es) registrada(s) en esta sesión**")
        st.dataframe(df_registro, use_container_width=True, hide_index=True)

        # ── Gráfico comparativo de escenarios ─────────────────────────────
        if len(df_registro) > 1:
            st.markdown("#### 📊 Comparativo de Budget Ajustado LRP por escenario")
            fy_labels_r = [f"FY{str(datetime.datetime.now().year + i)[2:]}" for i in range(1, 6)]
            cols_ajust = [f"Ajustado {fy} (US$ MM)" for fy in fy_labels_r]
            cols_disponibles = [c for c in cols_ajust if c in df_registro.columns]

            if cols_disponibles:
                comp_largo = df_registro[["Escenario"] + cols_disponibles].melt(
                    id_vars="Escenario", var_name="Año", value_name="Budget Ajustado (US$ MM)"
                )
                comp_largo["Año"] = comp_largo["Año"].str.extract(r"(FY\d+)")
                fig_comp = px.line(
                    comp_largo, x="Año", y="Budget Ajustado (US$ MM)", color="Escenario",
                    markers=True, title="Evolución LRP por escenario de sensibilidad",
                )
                st.plotly_chart(fig_comp, use_container_width=True)

            st.markdown("#### 🔢 Impacto total por escenario (US$ MM)")
            fig_bar_comp = px.bar(
                df_registro, x="Escenario", y="Impacto Sens. (US$ MM)",
                color="Impacto Sens. (US$ MM)",
                color_continuous_scale=["#4C78A8", "#FFFFFF", "#E45756"],
                color_continuous_midpoint=0,
                title="Impacto neto de sensibilidades por escenario (US$ MM)",
                text="Impacto Sens. (US$ MM)",
            )
            fig_bar_comp.update_traces(texttemplate="%{text:.2f}", textposition="outside")
            fig_bar_comp.update_layout(showlegend=False, height=380)
            st.plotly_chart(fig_bar_comp, use_container_width=True)

        # ── Descarga Excel del registro ────────────────────────────────────
        def registro_a_excel(df_reg: pd.DataFrame) -> bytes:
            out = io.BytesIO()
            with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
                wb = writer.book
                header_fmt = wb.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1})
                money_fmt  = wb.add_format({"num_format": '#,##0.00'})

                df_reg.to_excel(writer, index=False, sheet_name="Registro_Sensibilidades")
                ws = writer.sheets["Registro_Sensibilidades"]
                for c, col in enumerate(df_reg.columns):
                    ws.write(0, c, col, header_fmt)
                    width = min(max(len(str(col)) + 2, 14), 38)
                    ws.set_column(c, c, width)

                # Hoja resumen por año
                fy_labels_e = [f"FY{str(datetime.datetime.now().year + i)[2:]}" for i in range(1, 6)]
                cols_base_e    = [f"Ajustado {fy} (US$ MM)" for fy in fy_labels_e if f"Ajustado {fy} (US$ MM)" in df_reg.columns]
                cols_impacto_e = [f"Impacto {fy} (US$ MM)"  for fy in fy_labels_e if f"Impacto {fy} (US$ MM)"  in df_reg.columns]

                if cols_base_e:
                    resumen_anual = df_reg[["Escenario", "Timestamp"] + cols_base_e + cols_impacto_e]
                    resumen_anual.to_excel(writer, index=False, sheet_name="Resumen_Anual_Escenarios")
                    ws2 = writer.sheets["Resumen_Anual_Escenarios"]
                    for c, col in enumerate(resumen_anual.columns):
                        ws2.write(0, c, col, header_fmt)
                        ws2.set_column(c, c, 22)

            return out.getvalue()

        excel_registro = registro_a_excel(df_registro)
        st.download_button(
            label="⬇️ Descargar Registro completo de Sensibilidades (Excel)",
            data=excel_registro,
            file_name=f"registro_sensibilidades_lrp_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

        # ── Limpiar registro ───────────────────────────────────────────────
        if st.button("🗑️ Limpiar registro de esta sesión", use_container_width=True):
            st.session_state["registro_sensibilidades"] = []
            st.rerun()

# ─────────────────────────────────────────────
# PESTAÑA 5: INFORME EJECUTIVO
# ─────────────────────────────────────────────
with tab_informe:
    st.subheader("📄 Informe Ejecutivo")
    st.caption("Genera y descarga el informe completo con KPIs, tablas detalladas, simulación LRP y registro de sensibilidades.")

    if len(filtered) == 0:
        st.warning("⚠️ No hay datos cargados. Sube un archivo Excel y ajusta los filtros.")
    else:
        # Necesitamos resumen_5y para el informe — recalcularlo si no existe en scope
        año_base_inf = datetime.datetime.now().year
        df_5y_inf = simular_5_anos(filtered, cagr_inf, cagr_ops)
        cols_agrup_inf = ["Contexto_Mina"]
        if "Naturaleza" in df_5y_inf.columns:
            cols_agrup_inf = ["Contexto_Mina", "Naturaleza"]
        año_cols_inf = ["Año_0_Base", "Año_1", "Año_2", "Año_3", "Año_4", "Año_5"]
        resumen_5y_inf_raw = df_5y_inf.groupby(cols_agrup_inf)[año_cols_inf].sum().reset_index()
        estac_inf = calcular_estacionalidad_mensual(filtered, actual_months, forecast_months)
        año1_lbl_inf = año_base_inf + 1
        meses_año1_inf = [f"{m}-{str(año1_lbl_inf)[2:]}" for m in MONTH_ORDER]
        for m, mc in zip(MONTH_ORDER, meses_año1_inf):
            resumen_5y_inf_raw[mc] = resumen_5y_inf_raw["Año_1"] * estac_inf[m]
        rename_inf = {
            "Año_0_Base": f"FY{str(año_base_inf)[2:]}",
            "Año_2": f"FY{str(año_base_inf+2)[2:]}",
            "Año_3": f"FY{str(año_base_inf+3)[2:]}",
            "Año_4": f"FY{str(año_base_inf+4)[2:]}",
            "Año_5": f"FY{str(año_base_inf+5)[2:]}",
        }
        resumen_5y_inf = resumen_5y_inf_raw.rename(columns=rename_inf).drop(columns=["Año_1"])
        fy_ac_inf = [f"FY{str(año_base_inf+i)[2:]}" for i in [0, 2, 3, 4, 5]]
        cols_ord_inf = cols_agrup_inf + [fy_ac_inf[0]] + meses_año1_inf + fy_ac_inf[1:]
        resumen_5y_inf = resumen_5y_inf[[c for c in cols_ord_inf if c in resumen_5y_inf.columns]]

        st.markdown("---")
        st.markdown("### 🌐 Idioma / Language")
        col_lang1, col_lang2 = st.columns(2)
        gen_es = col_lang1.checkbox("🇨🇱 Español", value=True)
        gen_en = col_lang2.checkbox("🇺🇸 English", value=True)

        st.markdown("### 📂 Formato")
        col_fmt1, col_fmt2 = st.columns(2)
        gen_pdf  = col_fmt1.checkbox("PDF", value=True)
        gen_docx = col_fmt2.checkbox("Word (.docx)", value=True)

        st.markdown("---")

        # Preview de lo que incluirá el informe
        with st.expander("📋 Vista previa del contenido del informe", expanded=True):
            p1, p2, p3, p4 = st.columns(4)
            p1.metric("Actual YTD", money(actual_total))
            p2.metric("Budget FY",  money(budget_total))
            p3.metric("Forecast FY", money(forecast_total))
            p4.metric("Variación",  f"{var_pct:+.1%}")
            st.markdown(f"""
**El informe incluirá:**
- ✅ KPIs ejecutivos (Actual YTD, Budget FY, Forecast FY, Variación)
- ✅ Parámetros del modelo (sensibilidades, inflación, crecimiento)
- ✅ Tabla resumen por categoría ({len(summary)} categorías)
- ✅ Simulación LRP 5 años ({año1_lbl_inf}–{año_base_inf+5})
- ✅ Hallazgos automáticos y recomendaciones
- ✅ Registro de sensibilidades ({len(st.session_state['registro_sensibilidades'])} escenario(s) guardado(s))
            """)

        st.markdown("---")

        if st.button("🚀 Generar Informe(s)", use_container_width=True, type="primary"):
            args_informe = dict(
                actual_total=actual_total,
                budget_total=budget_total,
                forecast_total=forecast_total,
                var_total=var_total,
                var_pct=var_pct,
                summary=summary,
                chart_dim=chart_dim,
                actual_months=actual_months,
                forecast_months=forecast_months,
                cutoff_month=cutoff_month,
                sensitivity_mult=sensitivity_mult,
                momentum_mult=momentum_mult,
                cagr_inf=cagr_inf,
                cagr_ops=cagr_ops,
                resumen_5y=resumen_5y_inf,
                registro=st.session_state["registro_sensibilidades"],
            )

            langs = []
            if gen_es: langs.append("es")
            if gen_en: langs.append("en")

            if not langs:
                st.warning("Selecciona al menos un idioma.")
            elif not gen_pdf and not gen_docx:
                st.warning("Selecciona al menos un formato.")
            else:
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M")
                generated = {}

                with st.spinner("Generando informe(s)..."):
                    for lang in langs:
                        lang_label = "ES" if lang == "es" else "EN"
                        if gen_pdf:
                            if not REPORTLAB_OK:
                                st.error("❌ ReportLab no instalado. Agrega 'reportlab' a requirements.txt y redespliega.")
                            else:
                                try:
                                    pdf_bytes = generar_pdf_informe(lang=lang, **args_informe)
                                    generated[f"pdf_{lang}"] = (pdf_bytes, f"informe_forecast_{lang_label}_{ts}.pdf",
                                                                 "application/pdf")
                                except Exception as e:
                                    st.error(f"Error PDF {lang_label}: {e}")
                        if gen_docx:
                            if not PYTHON_DOCX_OK:
                                st.error("❌ python-docx no instalado. Agrega 'python-docx' a requirements.txt y redespliega.")
                            else:
                                try:
                                    docx_bytes = generar_docx_informe(lang=lang, **args_informe)
                                    generated[f"docx_{lang}"] = (docx_bytes, f"informe_forecast_{lang_label}_{ts}.docx",
                                                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                except Exception as e:
                                    st.error(f"Error DOCX {lang_label}: {e}")

                if generated:
                    st.success(f"✅ {len(generated)} archivo(s) listo(s). Descárgalos a continuación:")
                    for key, (data, fname, mime) in generated.items():
                        lang_flag = "🇨🇱" if "_es" in key else "🇺🇸"
                        fmt_icon  = "📄" if "pdf" in key else "📝"
                        fmt_label = "PDF" if "pdf" in key else "Word"
                        st.download_button(
                            label=f"{lang_flag} {fmt_icon} Descargar {fmt_label} — {'Español' if '_es' in key else 'English'}",
                            data=data,
                            file_name=fname,
                            mime=mime,
                            use_container_width=True,
                            key=f"dl_{key}_{ts}",
                        )

st.caption("Modelo diseñado para defensa académica: combina ejecución acumulada, tendencia reciente, curva no lineal y contexto minero.")
